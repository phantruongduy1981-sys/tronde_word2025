import streamlit as st
import re
import random
import zipfile
import io
import pandas as pd
from xml.dom import minidom
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==================== CẤU HÌNH GIAO DIỆN ====================
st.set_page_config(
    page_title="Trộn Đề THPT Minh Đức",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="collapsed"
)

HEADER_COLOR = "#00695c"
BUTTON_COLOR = "#d32f2f"
BG_INFO = "#e8f5e9"
BG_WARNING = "#fffde7"
CARD_BG = "#f1f8e9"

st.markdown(f"""
<style>
    .main-header {{ background-color: {HEADER_COLOR}; color: white; padding: 15px; text-align: center; border-radius: 8px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
    .main-header h1 {{ font-size: 22px; font-weight: 800; margin: 0; text-transform: uppercase; letter-spacing: 1px; }}
    .main-header p {{ margin-top: 5px; font-size: 13px; opacity: 0.9; font-weight: 500; }}
    .config-card {{ background-color: {CARD_BG}; border: 1px solid #c5e1a5; border-radius: 8px; padding: 15px; text-align: center; margin-bottom: 10px; }}
    .config-label {{ color: {HEADER_COLOR}; font-weight: bold; font-size: 14px; margin-bottom: 5px; display: block; }}
    .stButton>button {{ background-color: {BUTTON_COLOR}; color: white; border: none; border-radius: 8px; font-weight: bold; width: 100%; height: 50px; font-size: 16px; transition: 0.3s; }}
    .stButton>button:hover {{ background-color: #b71c1c; box-shadow: 0 4px 8px rgba(0,0,0,0.2); }}
    .info-box {{ background-color: {BG_INFO}; border-left: 5px solid #66bb6a; padding: 15px; border-radius: 5px; margin-bottom: 15px; }}
    .warning-box {{ background-color: {BG_WARNING}; border-left: 5px solid #fbc02d; padding: 15px; border-radius: 5px; margin-top: 15px; }}
    .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #eee; color: #888; font-size: 13px; }}
</style>
""", unsafe_allow_html=True)

# ==================== CORE LOGIC ====================
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def get_text_from_node(node):
    """Lấy text từ node"""
    texts = []
    for t in node.getElementsByTagNameNS(W_NS, "t"):
        if t.firstChild: texts.append(t.firstChild.nodeValue)
    return "".join(texts).strip()

def set_text_to_node(node, new_text):
    """Gán text mới vào node (xóa các t cũ, tạo 1 t mới)"""
    # Tìm run đầu tiên
    runs = node.getElementsByTagNameNS(W_NS, "r")
    if not runs: return
    run = runs[0]
    
    # Xóa hết t cũ trong run đó
    for t in run.getElementsByTagNameNS(W_NS, "t"):
        run.removeChild(t)
        
    # Tạo t mới
    doc = node.ownerDocument
    new_t = doc.createElementNS(W_NS, "w:t")
    # Thuộc tính preserve space
    new_t.setAttribute("xml:space", "preserve")
    text_node = doc.createTextNode(new_text)
    new_t.appendChild(text_node)
    run.appendChild(new_t)
    
    # Xóa các run thừa phía sau (để tránh text cũ còn sót)
    for i in range(1, len(runs)):
        node.removeChild(runs[i])

def check_structure_errors(blocks):
    full_text = "\n".join([get_text_from_node(b) for b in blocks])
    errors = []
    if not re.search(r'Câu\s*1[\.:]', full_text, re.IGNORECASE):
        errors.append("❌ Lỗi: Không tìm thấy 'Câu 1'. File phải bắt đầu bằng Câu 1.")
    return errors

def is_marked_correct(paragraph):
    """Kiểm tra đáp án đúng (Màu đỏ hoặc Gạch chân)"""
    runs = paragraph.getElementsByTagNameNS(W_NS, "r")
    for run in runs:
        rPr = run.getElementsByTagNameNS(W_NS, "rPr")
        if not rPr: continue
        rPr = rPr[0]
        colors = rPr.getElementsByTagNameNS(W_NS, "color")
        for c in colors:
            val = c.getAttributeNS(W_NS, "val")
            if val and (val.upper() in ['FF0000', 'RED']): return True
        u_tags = rPr.getElementsByTagNameNS(W_NS, "u")
        for u in u_tags:
            val = u.getAttributeNS(W_NS, "val")
            if val and val != 'none': return True
    return False

def clean_formatting(paragraph):
    """Xóa định dạng đỏ/gạch chân"""
    runs = paragraph.getElementsByTagNameNS(W_NS, "r")
    for run in runs:
        rPr_list = run.getElementsByTagNameNS(W_NS, "rPr")
        if not rPr_list: continue
        rPr = rPr_list[0]
        for c in rPr.getElementsByTagNameNS(W_NS, "color"): rPr.removeChild(c)
        for u in rPr.getElementsByTagNameNS(W_NS, "u"): rPr.removeChild(u)

def update_label_in_node(paragraph, new_label):
    """Cập nhật nhãn (Deep Scan)"""
    t_nodes = paragraph.getElementsByTagNameNS(W_NS, "t")
    if not t_nodes: return
    
    # Gom toàn bộ text lại để regex
    full_text = ""
    for t in t_nodes: 
        if t.firstChild: full_text += t.firstChild.nodeValue
            
    # Regex tìm vị trí nhãn cũ
    # Tìm A. B. C. D. hoặc a) b) c) d) hoặc Câu X
    pattern = r'^(\s*)([A-Da-d])([\.:\)])(.*)'
    match = re.match(pattern, full_text)
    
    if match:
        # Nếu match dạng A. hoặc a)
        prefix_space = match.group(1)
        # old_char = match.group(2)
        # suffix_char = match.group(3)
        rest = match.group(4)
        
        final_text = f"{prefix_space}{new_label}{rest}"
        set_text_to_node(paragraph, final_text)
        return

    # Check dạng Câu 1.
    match_q = re.match(r'^(\s*)(Câu\s*\d+)([\.:]*)?(.*)', full_text, re.IGNORECASE)
    if match_q:
        prefix = match_q.group(1)
        rest = match_q.group(4)
        final_text = f"{prefix}{new_label}{rest}"
        set_text_to_node(paragraph, final_text)

def split_and_normalize_options(question_blocks):
    """
    Tách các dòng chứa nhiều đáp án (A. ... B. ...) thành các dòng riêng.
    Trả về list blocks mới đã chuẩn hóa.
    """
    normalized_blocks = []
    
    for block in question_blocks:
        txt = get_text_from_node(block)
        
        # Regex tìm A., B., C., D. nằm giữa dòng
        # Pattern: Tìm vị trí bắt đầu của A., B., C., D.
        # Lưu ý: Cần cẩn thận với số thập phân (ví dụ 1.5) -> bắt buộc sau chấm phải là khoảng trắng hoặc số
        
        # Logic đơn giản: Nếu dòng chứa cả "A." và "B." thì tách
        # Tạm thời hỗ trợ tách cơ bản dựa trên "X. "
        matches = list(re.finditer(r'(?:\s|^)([A-D])[\.:]\s', txt))
        
        if len(matches) > 1:
            # Có nhiều đáp án trên 1 dòng -> Cần tách
            # Tạo các bản sao của block
            last_idx = 0
            
            # Cần check xem đáp án đúng nằm ở đâu trong dòng này (nếu có tô đỏ)
            # Việc này rất khó với XML node.
            # => GIẢI PHÁP: Nếu tách dòng, ta sẽ mất thuộc tính tô đỏ của từng đoạn con.
            # => HẠN CHẾ: Nếu dùng dòng gộp, đáp án đúng phải được nhận diện qua TEXT (ví dụ gạch chân text).
            # Ở đây để an toàn và đơn giản: Ta clone block, và set text lại.
            
            # Tuy nhiên, để giữ format (công thức toán), clone là tốt nhất.
            # Nhưng minidom không hỗ trợ cắt node con tương ứng với text range.
            
            # => GIẢI PHÁP THỰC TẾ CHO DỰ ÁN NÀY:
            # Nếu phát hiện dòng gộp, ta dùng Regex cắt text, tạo 4 block mới chỉ chứa Text (mất format ảnh/công thức trong dòng đó).
            # Đây là sự đánh đổi. Nếu muốn giữ format, người dùng nên xuống dòng.
            
            parts = []
            # Cắt chuỗi
            # Thêm sentinel vào cuối
            idxs = [m.start() for m in matches]
            idxs.append(len(txt))
            
            for i in range(len(matches)):
                start = idxs[i]
                end = idxs[i+1]
                sub_text = txt[start:end].strip()
                
                # Clone block gốc để lấy style
                new_blk = block.cloneNode(True)
                set_text_to_node(new_blk, sub_text)
                
                # Cố gắng detect lại đáp án đúng từ block gốc? 
                # Nếu block gốc tô đỏ toàn bộ -> tất cả con đều tô đỏ -> Sai.
                # => Tạm chấp nhận: Nếu dòng gộp, App sẽ coi như chưa tìm thấy đáp án đúng (User phải check Excel thủ công hoặc sửa file gốc xuống dòng).
                
                normalized_blocks.append(new_blk)
        else:
            normalized_blocks.append(block)
            
    return normalized_blocks

def shuffle_part1_mcq(question_blocks, q_idx):
    """Xử lý Phần 1: Trắc nghiệm"""
    
    # 1. Chuẩn hóa (Tách dòng gộp)
    normalized_blocks = split_and_normalize_options(question_blocks)
    
    intro = []
    options = []
    
    for b in normalized_blocks:
        txt = get_text_from_node(b)
        if re.match(r'^\s*[A-D][\.:]', txt):
            is_correct = is_marked_correct(b)
            options.append({'node': b, 'correct': is_correct, 'text': txt})
        else:
            intro.append(b)
    
    labels_mcq = ["A.", "B.", "C.", "D."]
    correct_char = "X"
    
    # Chỉ trộn nếu tìm thấy từ 2 đáp án trở lên (để tránh lỗi)
    if len(options) >= 2:
        random.shuffle(options)
        for i, opt in enumerate(options):
            lbl = labels_mcq[i] if i < 4 else "*"
            clean_formatting(opt['node']) # Xóa dấu hiệu
            update_label_in_node(opt['node'], lbl)
            if opt['correct']: correct_char = lbl[0]
        
        result_blocks = intro + [o['node'] for o in options]
    else:
        # Fallback: Không trộn, giữ nguyên
        for opt in options:
            if opt['correct']:
                m = re.match(r'^\s*([A-D])', opt['text'], re.IGNORECASE)
                if m: correct_char = m.group(1).upper()
            clean_formatting(opt['node'])
        result_blocks = intro + [o['node'] for o in options]

    if intro: 
        update_label_in_node(intro[0], f"Câu {q_idx}. ")
        
    return result_blocks, correct_char

def extract_part3_answer(question_blocks):
    """Phần 3: Trả lời ngắn"""
    answer_text = "X"
    blocks_to_keep = []
    found_ans = False
    for block in reversed(question_blocks):
        txt = get_text_from_node(block)
        match = re.search(r'(?:ĐS|DS)[:\s]+(.*)', txt, re.IGNORECASE)
        if match and not found_ans:
            if is_marked_correct(block):
                answer_text = match.group(1).strip()
                found_ans = True
                continue 
        blocks_to_keep.insert(0, block)
    return blocks_to_keep, answer_text

def shuffle_questions_router(questions, mode="MCQ", start_idx=1):
    processed_nodes = []
    key_map = {}
    
    for i, q_blocks in enumerate(questions):
        current_idx = start_idx + i
        
        if mode == "FILL":
            cleaned_blocks, ans_text = extract_part3_answer(q_blocks)
            if cleaned_blocks: 
                update_label_in_node(cleaned_blocks[0], f"Câu {current_idx}. ")
            processed_nodes.extend(cleaned_blocks)
            key_map[current_idx] = ans_text
            
        elif mode == "TF":
            # Phần 2: Giữ d cố định
            intro = []; options = []
            for b in q_blocks:
                txt = get_text_from_node(b)
                if re.match(r'^\s*[a-d][\.:\)]', txt, re.IGNORECASE):
                    is_c = is_marked_correct(b)
                    clean_formatting(b)
                    options.append({'node': b, 'text': txt})
                else: intro.append(b)
            
            d_node = None
            others = []
            for o in options:
                # Nhận diện d)
                if re.match(r'^\s*d[\.:\)]', o['text'], re.IGNORECASE): d_node = o
                else: others.append(o)
            
            random.shuffle(others)
            final_opts = others + ([d_node] if d_node else [])
            
            labels_tf = ["a)", "b)", "c)", "d)"]
            for k, opt in enumerate(final_opts):
                lbl = labels_tf[k] if k < 4 else "*"
                update_label_in_node(opt['node'], lbl)
            
            if intro: update_label_in_node(intro[0], f"Câu {current_idx}. ")
            processed_nodes.extend(intro + [o['node'] for o in final_opts])
            key_map[current_idx] = "X"
            
        else: # MCQ
            nodes, ans = shuffle_part1_mcq(q_blocks, current_idx)
            processed_nodes.extend(nodes)
            key_map[current_idx] = ans
            
    return processed_nodes, key_map

# --- TẠO FILE ĐÁP ÁN WORD ---
def create_word_answer_key(excel_data_list):
    doc = Document()
    doc.add_heading('BẢNG ĐÁP ÁN CHI TIẾT', 0)
    for data in excel_data_list:
        exam_code = data.get("Mã đề", "")
        doc.add_heading(f'Mã đề: {exam_code}', level=1)
        questions = {int(k): v for k, v in data.items() if k.isdigit()}
        sorted_q = sorted(questions.items())
        if not sorted_q: continue
        
        table = doc.add_table(rows=1, cols=10)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i in range(10): table.rows[0].cells[i].text = "Câu-ĐA"
        
        row_cells = None
        for idx, (q_num, ans) in enumerate(sorted_q):
            col = idx % 10
            if col == 0: row_cells = table.add_row().cells
            row_cells[col].text = f"{q_num}-{ans}"
        doc.add_paragraph("\n")
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

def process_docx(file_bytes, num_exams, start_id, shuffle_mode):
    input_io = io.BytesIO(file_bytes)
    resources = {}; xml_content = ""
    try:
        with zipfile.ZipFile(input_io, 'r') as zin:
            for filename in zin.namelist():
                if filename == "word/document.xml": xml_content = zin.read(filename).decode('utf-8')
                else: resources[filename] = zin.read(filename)
    except Exception as e: return None, None, None, [f"Lỗi đọc file: {str(e)}"]

    if not xml_content: return None, None, None, ["Không tìm thấy document.xml"]

    dom = minidom.parseString(xml_content)
    body = dom.getElementsByTagNameNS(W_NS, "body")[0]
    all_blocks = [child for child in list(body.childNodes) 
                  if child.nodeType == child.ELEMENT_NODE and child.localName in ["p", "tbl"]]
    errors = check_structure_errors(all_blocks)
    
    parts = []; current_part = []
    for block in all_blocks:
        txt = get_text_from_node(block)
        if re.match(r'^\s*PHẦN\s*\d+', txt, re.IGNORECASE):
            if current_part: parts.append(current_part)
            current_part = [block]
        else: current_part.append(block)
    if current_part: parts.append(current_part)
    if not parts: parts = [all_blocks]

    final_zip_io = io.BytesIO()
    excel_data_list = []
    
    with zipfile.ZipFile(final_zip_io, 'w', zipfile.ZIP_DEFLATED) as master_zip:
        for ver in range(num_exams):
            current_code = str(start_id + ver)
            curr_dom = minidom.parseString(xml_content)
            curr_body = curr_dom.getElementsByTagNameNS(W_NS, "body")[0]
            while curr_body.firstChild: curr_body.removeChild(curr_body.firstChild)
            
            exam_blocks = []; exam_key = {"Mã đề": current_code}; global_q_idx = 1
            
            for part_blocks in parts:
                questions = []; intro_part = []; curr_q = []; is_q = False
                cloned_blocks = [b.cloneNode(True) for b in part_blocks]
                
                for b in cloned_blocks:
                    txt = get_text_from_node(b)
                    if re.match(r'^\s*Câu\s*\d+', txt, re.IGNORECASE):
                        if curr_q: questions.append(curr_q)
                        curr_q = [b]; is_q = True
                    elif re.match(r'^\s*PHẦN', txt, re.IGNORECASE):
                        if curr_q: questions.append(curr_q)
                        curr_q = []; intro_part.append(b); is_q = False
                    else:
                        if is_q: curr_q.append(b)
                        else: intro_part.append(b)
                if curr_q: questions.append(curr_q)
                
                part_txt = get_text_from_node(intro_part[0]) if intro_part else ""
                current_mode = "MCQ"
                if shuffle_mode == "auto":
                    if "PHẦN 2" in part_txt.upper(): current_mode = "TF"
                    elif "PHẦN 3" in part_txt.upper(): current_mode = "FILL"
                elif shuffle_mode == "tf": current_mode = "TF"
                
                nodes, k_map = shuffle_questions_router(questions, current_mode, global_q_idx)
                for k, v in k_map.items(): exam_key[str(k)] = v
                global_q_idx += len(k_map)
                exam_blocks.extend(intro_part + nodes)

            for b in exam_blocks: curr_body.appendChild(b)
            
            new_xml = curr_dom.toxml()
            sub_io = io.BytesIO()
            with zipfile.ZipFile(sub_io, 'w', zipfile.ZIP_DEFLATED) as sub_z:
                sub_z.writestr("word/document.xml", new_xml.encode('utf-8'))
                for name, content in resources.items(): sub_z.writestr(name, content)
            
            master_zip.writestr(f"De_Thi/De_{current_code}.docx", sub_io.getvalue())
            excel_data_list.append(exam_key)
            
        if excel_data_list:
            df = pd.DataFrame(excel_data_list)
            cols = ["Mã đề"] + sorted([c for c in df.columns if c != "Mã đề"], key=lambda x: int(x) if x.isdigit() else 999)
            df = df[cols]
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False, sheet_name='DapAn')
            master_zip.writestr("Dap_An_Excel.xlsx", excel_io.getvalue())
            
            word_bytes = create_word_answer_key(excel_data_list)
            master_zip.writestr("Dap_An_Word.docx", word_bytes)
        
    return final_zip_io.getvalue(), None, None, errors

# ==================== MAIN UI ====================
def main():
    st.markdown("""
    <div class="main-header">
        <h1>TRƯỜNG TRUNG HỌC PHỔ THÔNG MINH ĐỨC</h1>
        <p>APP TRỘN ĐỀ 2025 - PRO VERSION 3.1 (Fix Core)</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        with st.expander("📄 Hướng dẫn & File Mẫu", expanded=True):
            sample_url = "https://docs.google.com/document/d/1i1b-By6EA_HO8fWgMYG9iXZPGannmWdg/export?format=docx"
            st.link_button("📥 Tải File Mẫu", sample_url, use_container_width=True)
            st.markdown("""
            <div class="info-box" style="margin-top:15px;">
                <div class="info-header">📌 Cấu trúc:</div>
                <ul>
                    <li><b>PHẦN 1:</b> Trắc nghiệm (A.B.C.D)</li>
                    <li><b>PHẦN 2:</b> Đúng/Sai (a, b, c, d) - <i>(d cố định)</i></li>
                    <li><b>PHẦN 3:</b> Trả lời ngắn</li>
                </ul>
            </div>
            <div class="warning-box">
                <div class="warning-header">⚠️ Lưu ý Quan Trọng:</div>
                <ul>
                    <li><b>Quy định chung:</b> Câu hỏi bắt đầu bằng <code>Câu 1.</code> (1 dấu chấm).</li>
                    <li><b>Đáp án đúng (P1, P2):</b> Phải <b>Gạch chân</b> hoặc <span style="color:red"><b>Tô đỏ</b></span>.</li>
                    <li><b>Đáp án P3:</b> Ghi <span style="color:red"><b>ĐS: Kết quả</b></span> và tô đỏ ở cuối câu.</li>
                    <li><b>Khuyên dùng:</b> Mỗi đáp án A, B, C, D nên xuống dòng để tránh lỗi hiển thị.</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('<div class="section-title"><span class="step-circle">1</span> Chọn file đề Word (*.docx)</div>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader("", type=["docx"], label_visibility="collapsed")
        if uploaded_file: st.success(f"✅ Đã tải: {uploaded_file.name}")
        
    with col2:
        st.markdown('<div class="section-title"><span class="step-circle">2</span> Cấu hình</div>', unsafe_allow_html=True)
        shuffle_opt = st.radio("", ["🔄 Tự động", "📝 Trắc nghiệm", "✅ Đúng/Sai"], index=0, label_visibility="collapsed")
        mode_map = {"🔄 Tự động": "auto", "📝 Trắc nghiệm": "mcq", "✅ Đúng/Sai": "tf"}
        selected_mode = mode_map[shuffle_opt]

        st.write("")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown('<div class="config-card"><span class="config-label">SỐ LƯỢNG ĐỀ</span></div>', unsafe_allow_html=True)
            num_exams = st.number_input("Số lượng", 1, 50, 4, label_visibility="collapsed")
        with c2:
            st.markdown('<div class="config-card"><span class="config-label">MÃ BẮT ĐẦU</span></div>', unsafe_allow_html=True)
            start_id = st.number_input("Mã bắt đầu", 1, 9999, 1001, label_visibility="collapsed")
            
        st.caption(f"📍 Sẽ tạo các mã đề từ: {start_id} ➝ {start_id + num_exams - 1}")
        st.markdown("<br>", unsafe_allow_html=True)

        if st.button("🚀 Trộn đề & Tải Trọn Bộ (ZIP)"):
            if not uploaded_file: st.error("⚠️ Vui lòng chọn file đề trước!")
            else:
                with st.spinner("⏳ Đang xử lý..."):
                    try:
                        file_bytes = uploaded_file.read()
                        uploaded_file.seek(0)
                        zip_data, _, _, errors = process_docx(file_bytes, num_exams, start_id, selected_mode)
                        
                        if errors:
                            for e in errors: st.error(e)
                        
                        if zip_data:
                            st.download_button("📦 Tải Trọn Bộ (ZIP)", zip_data, f"Tron_Bo_De_{uploaded_file.name}.zip", "application/zip", type="primary")
                            st.balloons()
                            
                    except Exception as e: st.error(f"Lỗi: {str(e)}")

    st.markdown("""<div class="footer">© 2025 Phan Trường Duy - THPT Minh Đức<br>PRO VERSION 3.1</div>""", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
